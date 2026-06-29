import docx
import os

base_dir = r"d:\6th Sem\Redrob AI\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge"

for fn in ["README.docx", "redrob_signals_doc.docx", "submission_spec.docx", "job_description.docx"]:
    print(f"=== {fn} ===")
    doc = docx.Document(os.path.join(base_dir, fn))
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(p.text)
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"Row {r_idx}: {cells}")
    print()
